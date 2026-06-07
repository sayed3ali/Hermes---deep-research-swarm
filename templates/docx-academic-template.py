#!/usr/bin/env python3
"""
Academic DOCX Template for deep-research-swarm
Based on: Transferable Skills-food.docx analysis
Style: Professional academic report with Cite Them Right Harvard referencing

Color palette (from example document):
- Primary dark: #1F3864 (Heading 1, table headers)
- Secondary blue: #2E75B6 (Part labels, accents)
- Heading 2: #2A4D7A
- Heading 3: #1F4D78
- Heading 4-6: #2E74B5
- Body text: #000000 (default)
- Subtitle: #595959
- Table borders: #BFBFBF
- Theme accent1: #4472C4

Font: Calibri (default Office font)
Page: US Letter 8.5" x 11", margins 1" all sides
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# === COLOR PALETTE ===
C_PRIMARY = RGBColor(0x1F, 0x38, 0x64)      # Dark navy - H1, table headers
C_SECONDARY = RGBColor(0x2E, 0x75, 0xB6)   # Medium blue - Part labels
C_H2 = RGBColor(0x2A, 0x4D, 0x7A)           # Heading 2
C_H3 = RGBColor(0x1F, 0x4D, 0x78)           # Heading 3
C_H4 = RGBColor(0x2E, 0x74, 0xB5)           # Heading 4-6
C_SUBTITLE = RGBColor(0x59, 0x59, 0x59)     # Gray - subtitles
C_TEXT = RGBColor(0x00, 0x00, 0x00)         # Black - body text
C_TABLE_BORDER = "BFBFBF"
C_TABLE_HEADER_BG = "1F3864"

# === FONT SETTINGS ===
FONT_NAME = "Calibri"
FONT_SIZE_BODY = 11  # pt
FONT_SIZE_H1 = 15
FONT_SIZE_H2 = 12.5
FONT_SIZE_H3 = 12
FONT_SIZE_H4 = 11
FONT_SIZE_TITLE = 24
FONT_SIZE_SUBTITLE = 12
FONT_SIZE_PART = 12
FONT_SIZE_FOOTNOTE = 10

# === PAGE SETTINGS ===
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11.0)
MARGIN = Inches(1.0)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, color=C_TABLE_BORDER, size="4", val="single"):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        edge_elem = OxmlElement(f'w:{edge}')
        edge_elem.set(qn('w:val'), val)
        edge_elem.set(qn('w:sz'), size)
        edge_elem.set(qn('w:color'), color)
        tcBorders.append(edge_elem)
    tcPr.append(tcBorders)


def set_run_font(run, name=FONT_NAME, size=FONT_SIZE_BODY, bold=False, italic=False, color=C_TEXT):
    """Apply font formatting to a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_hanging_indent(paragraph, indent_inches=0.33):
    """Add hanging indent for references."""
    pf = paragraph.paragraph_format
    pf.left_indent = Inches(indent_inches)
    pf.first_line_indent = Inches(-indent_inches)


def create_academic_docx(title, subtitle, part_label=None):
    """Create a new academic-style document matching the example format."""
    doc = Document()

    # Set page size and margins
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN

    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE_BODY)
    font.color.rgb = C_TEXT

    # Configure heading styles
    for level, (size, color, bold) in {
        1: (FONT_SIZE_H1, C_PRIMARY, True),
        2: (FONT_SIZE_H2, C_H2, True),
        3: (FONT_SIZE_H3, C_H3, False),
        4: (FONT_SIZE_H4, C_H4, False),
    }.items():
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = FONT_NAME
        h_style.font.size = Pt(size)
        h_style.font.bold = bold
        h_style.font.color.rgb = color
        h_style.paragraph_format.space_after = Pt(6)
        h_style.paragraph_format.space_before = Pt(12)

    # === COVER / TITLE PAGE ===
    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    set_run_font(run, size=FONT_SIZE_TITLE, bold=True, color=C_PRIMARY)
    p.paragraph_format.space_after = Pt(6)

    # Subtitle
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(subtitle)
        set_run_font(run, size=FONT_SIZE_SUBTITLE, italic=True, color=C_SUBTITLE)
        p.paragraph_format.space_after = Pt(30)

    # Part label (e.g., "Part 1")
    if part_label:
        p = doc.add_paragraph()
        run = p.add_run(part_label)
        set_run_font(run, size=FONT_SIZE_PART, bold=True, color=C_SECONDARY)
        p.paragraph_format.space_after = Pt(3)

    return doc


def add_body_paragraph(doc, text, bold_start=None):
    """Add a justified body paragraph with optional bold start.

    Args:
        doc: Document object
        text: Full paragraph text
        bold_start: Text at the start that should be bold (e.g., "Resilience, Flexibility")
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(10)

    if bold_start and text.startswith(bold_start):
        run = p.add_run(bold_start)
        set_run_font(run, bold=True)
        remainder = text[len(bold_start):]
        if remainder:
            run = p.add_run(remainder)
            set_run_font(run)
    else:
        run = p.add_run(text)
        set_run_font(run)

    return p


def add_reference(doc, text):
    """Add a reference with hanging indent (Cite Them Right Harvard format).

    Format examples:
    - Book: Author, A. (Year) Title. Place: Publisher.
    - Journal: Author, A. (Year) 'Title', Journal Name, Volume(Issue), pp. pages.
    - Web: Author. (Year) Title. Available at: URL (Accessed: Date).
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run)
    add_hanging_indent(p, indent_inches=0.33)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_info_table(doc, data):
    """Add an information table (like the student info table in the example).

    Args:
        data: List of [label, value] pairs
    """
    table = doc.add_table(rows=len(data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (label, value) in enumerate(data):
        row = table.rows[i]
        # Label cell
        cell_label = row.cells[0]
        cell_label.text = label
        set_run_font(cell_label.paragraphs[0].runs[0], bold=True)
        set_cell_border(cell_label)

        # Value cell
        cell_value = row.cells[1]
        cell_value.text = value
        set_cell_border(cell_value)

    doc.add_paragraph()  # Space after table
    return table


def add_data_table(doc, headers, rows):
    """Add a data table with styled header row.

    Args:
        headers: List of column header strings
        rows: List of row data (each row is a list of cell values)
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_shading(cell, C_TABLE_HEADER_BG)
        set_cell_border(cell)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(value)
            set_cell_border(cell)
            for run in cell.paragraphs[0].runs:
                set_run_font(run)

    doc.add_paragraph()  # Space after table
    return table


def add_footer_with_page_number(doc, text_left=""):
    """Add a footer with left text and right page number.

    Args:
        text_left: Text to display on the left (e.g., "Critical Review")
    """
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Left text
    if text_left:
        run = p.add_run(text_left)
        set_run_font(run, size=FONT_SIZE_FOOTNOTE, color=C_SUBTITLE)
        run = p.add_run("    |   ")
        set_run_font(run, size=FONT_SIZE_FOOTNOTE, color=C_SUBTITLE)

    # Page number field
    run = p.add_run()
    set_run_font(run, size=FONT_SIZE_FOOTNOTE, color=C_SUBTITLE)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = p.add_run()
    set_run_font(run2, size=FONT_SIZE_FOOTNOTE, color=C_SUBTITLE)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run2._r.append(instrText)

    run3 = p.add_run()
    set_run_font(run3, size=FONT_SIZE_FOOTNOTE, color=C_SUBTITLE)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = p.add_run("2")  # Placeholder
    set_run_font(run4, size=FONT_SIZE_FOOTNOTE, color=C_SUBTITLE)

    run5 = p.add_run()
    set_run_font(run5, size=FONT_SIZE_FOOTNOTE, color=C_SUBTITLE)
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)


# === CITE THEM RIGHT HARVARD REFERENCE FORMAT ===

HARVARD_FORMAT_GUIDE = """
Cite Them Right Harvard Reference Format (12th Edition)

In-text citations:
- One author: (Smith, 2020)
- Two authors: (Smith and Jones, 2019)
- Three+ authors: (Smith et al., 2018)
- Direct quote: (Smith, 2020, p. 45)
- No author: ('Title of Article', 2021) or (Organisation Name, 2020)
- No date: (Smith, no date)

Reference list format:

BOOK:
Author, A.A. (Year) Title of Book. Place of Publication: Publisher.
Example: Bryman, A. (2016) Social Research Methods. 5th edn. Oxford: Oxford University Press.

EDITED BOOK:
Editor, A.A. (ed.) (Year) Title of Book. Place: Publisher.

JOURNAL ARTICLE:
Author, A.A. (Year) 'Title of article', Title of Journal, Volume(Issue), pp. page range.
Example: Clegg, S.R., Redding, S.G. and Cartner, M. (1990) 'Organisation and management in East Asia', Organisation Studies, 11(1), pp. 123-145.

WEBSITE / ONLINE:
Author or Organisation. (Year) Title of webpage. Available at: URL (Accessed: Day Month Year).
Example: BBC News. (2021) Climate change: UK aims to cut emissions by 78% by 2035. Available at: https://www.bbc.co.uk/news (Accessed: 21 April 2021).

REPORT:
Author or Organisation. (Year) Title of Report. Place: Publisher.
Example: World Economic Forum. (2025) The Future of Jobs Report 2025. Geneva: World Economic Forum.

GOVERNMENT PUBLICATION:
Government Name. (Year) Title. Place: Publisher.
Example: UK Government. (2023) Policy Paper: AI Regulation. London: HM Government.

NEWSPAPER ARTICLE:
Author, A.A. (Year) 'Title of article', Title of Newspaper, Day Month, p. page.

CONFERENCE PAPER:
Author, A.A. (Year) 'Title of paper', in Title of Conference Proceedings. Place: Publisher, pp. pages.

DISSERTATION / THESIS:
Author, A.A. (Year) 'Title of dissertation', Degree level. Institution.

PERSONAL COMMUNICATION:
Author, A.A. (Year) Personal communication, Day Month.

Key rules:
1. Reference list is alphabetised by author surname
2. Use 'pp.' for page ranges, 'p.' for single page
3. Use single quotation marks for article titles
4. Use italics for book and journal titles
5. Include 'Available at:' and 'Accessed:' for online sources
6. For multiple works by same author, order by year (earliest first)
7. For same author + same year, use a, b, c suffixes: (Smith, 2020a), (Smith, 2020b)
"""


if __name__ == "__main__":
    # Example usage
    doc = create_academic_docx(
        title="Quantum Computing Commercial Landscape 2026",
        subtitle="A critical review of market dynamics, technical challenges, and competitive positioning",
        part_label="Research Report"
    )

    # Add info table
    add_info_table(doc, [
        ["Topic", "Quantum Computing Market Analysis"],
        ["Date", "June 2026"],
        ["Method", "Multi-Agent Deep Research"],
    ])

    # Add headings and body text
    doc.add_heading("Executive Summary", level=1)
    add_body_paragraph(doc,
        "The global quantum computing market in 2026 is valued at approximately 698.6 million to 2.04 billion "
        "dollars, with analyst projections ranging from 1.72 billion by 2033 to 18.33 billion by 2034. "
        "North America dominates with 43.6 percent market share, driven by two billion dollars in US government "
        "CHIPS Act funding to nine quantum companies including IBM and D-Wave.")

    add_body_paragraph(doc,
        "Enterprise adoption has reached a tipping point, with over 300 global companies actively deploying "
        "quantum technology. A June 2026 survey of large UK enterprises found that 65 percent are piloting "
        "or adopting quantum solutions.")

    doc.add_heading("Market Size and Growth", level=2)
    add_body_paragraph(doc,
        "Analyst estimates for the 2026 quantum computing market vary significantly, reflecting both rapid "
        "growth and methodological uncertainty. Persistence Market Research places the market at 698.6 million "
        "dollars in 2026, projecting growth to 1.72 billion by 2033 at a compound annual growth rate of 13.7 "
        "percent.",
        bold_start="Market Size and Growth. ")

    # Add data table
    add_data_table(doc,
        headers=["Region", "2026 Market Size", "2035 Projection", "CAGR"],
        rows=[
            ["North America", "$890M", "$4.2B", "18.7%"],
            ["Europe", "$530M", "$3.3B", "22.4%"],
            ["Asia-Pacific", "$420M", "$5.1B", "28.9%"],
            ["Rest of World", "$200M", "$1.1B", "20.8%"],
        ]
    )

    # Add references
    doc.add_page_break()
    doc.add_heading("Reference List", level=1)

    add_reference(doc,
        "McKinsey and Company. (2026) Quantum Technology Monitor 2026. Available at: "
        "https://www.mckinsey.com (Accessed: 7 June 2026).")

    add_reference(doc,
        "Fortune Business Insights. (2026) Quantum Computing Market Size, Share and Industry Analysis. "
        "Available at: https://www.fortunebusinessinsights.com (Accessed: 7 June 2026).")

    add_reference(doc,
        "World Economic Forum. (2025) The Future of Jobs Report 2025. Geneva: World Economic Forum.")

    # Add footer
    add_footer_with_page_number(doc, text_left="Research Report")

    # Save
    output_path = r"C:\Users\User\hermes\research\chart-test-2026-06-07\quantum-academic-style.docx"
    doc.save(output_path)
    print(f"Saved: {output_path}")
